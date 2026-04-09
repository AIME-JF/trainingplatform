"""
Word parser with optional python-docx dependency.
"""

from io import BytesIO
from typing import Any, List

from .base import BaseParser, DocumentMetadata, DocumentType, ParseResult
from .image_parser import ImageParser, ParseImage, ParseImageResult
from logger import logger


class WordParser(BaseParser):
    """Parse DOCX files and OCR embedded images."""

    def __init__(self) -> None:
        self.image_parser = ImageParser()

    @property
    def supported_extensions(self) -> list:
        return [".docx"]

    @property
    def parser_name(self) -> str:
        return "WordParser"

    def can_parse(self, filename: str, mime_type: str = None) -> bool:
        normalized_name = str(filename or "").lower()
        return normalized_name.endswith(tuple(self.supported_extensions)) or (
            mime_type and ("word" in mime_type.lower() or "document" in mime_type.lower())
        )

    def parse(self, file_data: bytes, filename: str) -> ParseResult:
        try:
            from docx import Document
        except ImportError as exc:
            logger.error("Word parser dependency missing for %s: %s", filename, exc)
            return self._empty_result(len(file_data))

        try:
            document = Document(BytesIO(file_data))
            images = self._extract_images(document)
            image_results = self.image_parser.process_images_parallel(images)
            document = self._insert_image_content(document, image_results)

            file_stream = BytesIO()
            document.save(file_stream)

            try:
                from markitdown import MarkItDown

                markdown_converter = MarkItDown()
                content = markdown_converter.convert(file_stream).text_content
            except ImportError:
                content = "\n".join(
                    paragraph.text for paragraph in document.paragraphs if paragraph.text.strip()
                )

            metadata = DocumentMetadata(
                parser=self.parser_name,
                document_type=DocumentType.WORD,
                size=len(file_data),
                lines=len(content.splitlines()),
            )
            return ParseResult(content, metadata)
        except Exception as exc:
            logger.error("Word parse failed for %s: %s", filename, exc)
            return self._empty_result(len(file_data))

    def _extract_images(self, document: Any) -> List[ParseImage]:
        images: List[ParseImage] = []
        for paragraph in document.paragraphs:
            for run in paragraph.runs:
                drawing_elements = run.element.xpath(".//w:drawing")
                if not drawing_elements:
                    continue
                blip_elements = run.element.xpath(".//a:blip")
                for blip in blip_elements:
                    embed_id = blip.get(
                        "{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed"
                    )
                    if embed_id and embed_id in document.part.related_parts:
                        image_data = document.part.related_parts[embed_id].blob
                        images.append(
                            ParseImage(
                                id=self.image_parser.get_image_hash(image_data),
                                index=len(images),
                                data=image_data,
                            )
                        )
                        break
        return images

    def _insert_image_content(
        self,
        document: Any,
        image_results: List[ParseImageResult],
    ) -> Any:
        for paragraph in document.paragraphs:
            for run in paragraph.runs:
                drawing_elements = run.element.xpath(".//w:drawing")
                if not drawing_elements:
                    continue
                blip_elements = run.element.xpath(".//a:blip")
                for blip in blip_elements:
                    embed_id = blip.get(
                        "{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed"
                    )
                    if embed_id and embed_id in document.part.related_parts:
                        image_data = document.part.related_parts[embed_id].blob
                        image_hash = self.image_parser.get_image_hash(image_data)
                        result = next((item for item in image_results if item.id == image_hash), None)
                        if result:
                            run.clear()
                            run.add_text(result.content)
                        break
        return document

    def _empty_result(self, file_size: int) -> ParseResult:
        return ParseResult(
            "",
            DocumentMetadata(
                parser=self.parser_name,
                document_type=DocumentType.WORD,
                size=file_size,
                lines=0,
            ),
        )
