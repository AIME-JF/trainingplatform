"""Document text extraction service — supports PDF, DOCX, XLSX, CSV, TXT."""
import csv
import io
from typing import Optional


class DocumentParserService:
    def parse(self, content: bytes, filename: str) -> str:
        ext = filename.rsplit(".", 1)[-1].lower() if "." in filename else ""
        parsers = {
            "pdf": self._parse_pdf,
            "docx": self._parse_docx,
            "doc": self._parse_docx,
            "xlsx": self._parse_xlsx,
            "xls": self._parse_xls,
            "csv": self._parse_csv,
            "txt": self._parse_txt,
            "md": self._parse_txt,
        }
        parser = parsers.get(ext)
        if not parser:
            raise ValueError(f"不支持的文件格式: .{ext}，支持 PDF/DOCX/XLSX/CSV/TXT")
        return parser(content)

    def _parse_pdf(self, content: bytes) -> str:
        try:
            from pypdf import PdfReader
        except ImportError:
            raise RuntimeError("缺少 pypdf 依赖，请运行: pip install pypdf")

        reader = PdfReader(io.BytesIO(content))
        parts = []
        for i, page in enumerate(reader.pages):
            text = page.extract_text()
            if text and text.strip():
                parts.append(text.strip())
        return "\n\n".join(parts) if parts else "[PDF 中未提取到文本内容]"

    def _parse_docx(self, content: bytes) -> str:
        try:
            from docx import Document
        except ImportError:
            raise RuntimeError("缺少 python-docx 依赖，请运行: pip install python-docx")

        doc = Document(io.BytesIO(content))
        parts = []
        for para in doc.paragraphs:
            if para.text.strip():
                parts.append(para.text.strip())
        for table in doc.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if cells:
                    parts.append(" | ".join(cells))
        return "\n".join(parts) if parts else "[DOCX 中未提取到文本内容]"

    def _parse_xlsx(self, content: bytes) -> str:
        try:
            from openpyxl import load_workbook
        except ImportError:
            raise RuntimeError("缺少 openpyxl 依赖")

        wb = load_workbook(filename=io.BytesIO(content), read_only=True, data_only=True)
        return self._read_workbook_sheets(wb)

    def _parse_xls(self, content: bytes) -> str:
        try:
            import xlrd
        except ImportError:
            raise RuntimeError("缺少 xlrd 依赖，请运行: pip install xlrd")

        wb = xlrd.open_workbook(file_contents=content)
        parts = []
        for sheet_name in wb.sheet_names():
            ws = wb.sheet_by_name(sheet_name)
            parts.append(f"=== 工作表: {sheet_name} ===")
            for row_idx in range(ws.nrows):
                cells = [str(ws.cell_value(row_idx, col_idx)).strip() for col_idx in range(ws.ncols)]
                cells = [c for c in cells if c]
                if cells:
                    parts.append(" | ".join(cells))
        return "\n".join(parts) if parts else "[XLS 中未提取到文本内容]"

    def _read_workbook_sheets(self, wb) -> str:
        parts = []
        for sheet_name in wb.sheetnames:
            ws = wb[sheet_name]
            parts.append(f"=== 工作表: {sheet_name} ===")
            for row in ws.iter_rows(values_only=True):
                cells = [str(c).strip() for c in row if c is not None and str(c).strip()]
                if cells:
                    parts.append(" | ".join(cells))
        return "\n".join(parts) if parts else "[XLSX 中未提取到文本内容]"

    def _parse_csv(self, content: bytes) -> str:
        for encoding in ("utf-8-sig", "utf-8", "gbk", "gb2312", "latin-1"):
            try:
                text = content.decode(encoding)
                break
            except (UnicodeDecodeError, LookupError):
                continue
        else:
            text = content.decode("utf-8", errors="replace")
        reader = csv.reader(io.StringIO(text))
        parts = []
        for row in reader:
            cells = [c.strip() for c in row if c.strip()]
            if cells:
                parts.append(" | ".join(cells))
        return "\n".join(parts) if parts else "[CSV 中未提取到文本内容]"

    def _parse_txt(self, content: bytes) -> str:
        for encoding in ("utf-8", "gbk", "gb2312", "latin-1"):
            try:
                return content.decode(encoding)
            except (UnicodeDecodeError, LookupError):
                continue
        return content.decode("utf-8", errors="replace")
